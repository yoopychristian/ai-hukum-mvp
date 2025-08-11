import os
import uuid
from datetime import datetime
from typing import Optional, List

import fitz  # PyMuPDF
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from sqlalchemy import create_engine, String, DateTime, Boolean, ForeignKey, Integer
from sqlalchemy.orm import sessionmaker, declarative_base, Mapped, mapped_column, relationship, Session
from fpdf import FPDF

try:
    from anthropic import Anthropic
except Exception as e:  # pragma: no cover
    Anthropic = None  # type: ignore


APP_NAME = "AI Hukum MVP Backend"

app = FastAPI(title=APP_NAME)

# CORS
cors_origins = os.getenv("CORS_ORIGINS", "*")
allow_all = cors_origins.strip() == "*"
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"] if allow_all else [o.strip() for o in cors_origins.split(",") if o.strip()],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# In-memory session storage
SESSIONS: dict[str, dict[str, str]] = {}


# SQLite (Level 2 features: chat history)
DB_URL = os.getenv("DATABASE_URL", "sqlite:///hukum.db")
engine = create_engine(DB_URL, connect_args={"check_same_thread": False} if DB_URL.startswith("sqlite") else {})
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()


class Chat(Base):
    __tablename__ = "chats"
    id: Mapped[str] = mapped_column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    title: Mapped[Optional[str]] = mapped_column(String, nullable=True)
    confidential: Mapped[bool] = mapped_column(Boolean, default=False)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    updated_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    messages: Mapped[List["Message"]] = relationship("Message", back_populates="chat", cascade="all, delete-orphan")
    files: Mapped[List["FileRec"]] = relationship("FileRec", back_populates="chat", cascade="all, delete-orphan")


class Message(Base):
    __tablename__ = "messages"
    id: Mapped[str] = mapped_column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    chat_id: Mapped[str] = mapped_column(String, ForeignKey("chats.id"))
    role: Mapped[str] = mapped_column(String)  # 'user' | 'assistant'
    content: Mapped[str] = mapped_column(String)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    chat: Mapped[Chat] = relationship("Chat", back_populates="messages")


class FileRec(Base):
    __tablename__ = "files"
    id: Mapped[str] = mapped_column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    chat_id: Mapped[str] = mapped_column(String, ForeignKey("chats.id"))
    name: Mapped[str] = mapped_column(String)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)
    chat: Mapped[Chat] = relationship("Chat", back_populates="files")


class Feedback(Base):
    __tablename__ = "feedback"
    id: Mapped[str] = mapped_column(String, primary_key=True, default=lambda: str(uuid.uuid4()))
    chat_id: Mapped[Optional[str]] = mapped_column(String, ForeignKey("chats.id"), nullable=True)
    message_id: Mapped[Optional[str]] = mapped_column(String, ForeignKey("messages.id"), nullable=True)
    value: Mapped[int] = mapped_column(Integer)  # 1 or -1
    comment: Mapped[Optional[str]] = mapped_column(String, nullable=True)
    created_at: Mapped[datetime] = mapped_column(DateTime, default=datetime.utcnow)


Base.metadata.create_all(bind=engine)


def get_db() -> Session:
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


class SummarizeRequest(BaseModel):
    session_id: str
    lang: Optional[str] = None  # 'id' or 'en'


class AskRequest(BaseModel):
    session_id: str
    question: str
    lang: Optional[str] = None  # 'id' or 'en'


class DraftRequest(BaseModel):
    session_id: Optional[str] = None
    doc_type: str
    requirements: Optional[str] = None
    tone: Optional[str] = None  # e.g., formal | neutral
    length: Optional[str] = None  # short | medium | long
    lang: Optional[str] = None  # 'id' or 'en'


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            texts: list[str] = []
            for page in doc:
                texts.append(page.get_text())
            return "\n".join(texts).strip()
    except Exception as e:  # pragma: no cover
        raise HTTPException(status_code=400, detail=f"Gagal membaca PDF: {e}")


def get_client() -> Anthropic:
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY belum diset di environment")
    if Anthropic is None:  # pragma: no cover
        raise HTTPException(status_code=500, detail="anthropic SDK tidak terinstal")
    return Anthropic(api_key=api_key)


def call_claude(prompt: str, max_tokens: int = 1024, temperature: float = 0.2) -> str:
    client = get_client()
    model = os.getenv("ANTHROPIC_MODEL", "claude-3-haiku-20240307")
    try:
        msg = client.messages.create(
            model=model,
            max_tokens=max_tokens,
            temperature=temperature,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                    ],
                }
            ],
        )
        # anthropic SDK returns a list of content blocks; we expect first to be text
        if hasattr(msg, "content") and len(msg.content) > 0:
            block = msg.content[0]
            # New SDK returns objects with .type and .text
            text = getattr(block, "text", None)
            if isinstance(text, str):
                return text
        return ""
    except Exception as e:  # pragma: no cover
        raise HTTPException(status_code=500, detail=f"Gagal memanggil model: {e}")


@app.get("/")
def root():
    return {"name": APP_NAME, "status": "ok"}


@app.post("/upload")
async def upload(
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
):
    parts: List[str] = []

    # single file
    if file is not None:
        content = await file.read()
        if not content:
            raise HTTPException(status_code=400, detail="File kosong")
        content_type = (file.content_type or "").lower()
        filename_lower = (file.filename or "").lower()
        is_pdf = "pdf" in content_type or filename_lower.endswith(".pdf")
        is_txt = "text" in content_type or filename_lower.endswith(".txt")
        if is_pdf:
            parts.append(extract_text_from_pdf(content))
        elif is_txt:
            try:
                parts.append(content.decode("utf-8", errors="ignore"))
            except Exception:
                parts.append(content.decode("latin-1", errors="ignore"))
        else:
            raise HTTPException(status_code=400, detail="Tipe file tidak didukung. Gunakan PDF atau TXT.")

    if text and text.strip():
        parts.append(text.strip())

    extracted_text = "\n\n---\n\n".join([p for p in parts if p]).strip()
    if not extracted_text:
        raise HTTPException(status_code=400, detail="Kirim file PDF/TXT atau teks pada field 'text'")

    session_id = str(uuid.uuid4())
    SESSIONS[session_id] = {"text": extracted_text}
    return {"session_id": session_id, "num_chars": len(extracted_text)}


@app.post("/summarize")
def summarize(req: SummarizeRequest):
    session = SESSIONS.get(req.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session tidak ditemukan")

    doc_text = session.get("text", "")
    if not doc_text:
        raise HTTPException(status_code=400, detail="Tidak ada teks dalam sesi")

    # Keep prompt within limits
    snippet = doc_text[:15000]
    language = (req.lang or "id").lower()
    # Make the target output language explicit so the model doesn't mirror the document's language
    lang_line = "Jawab ringkas dalam bahasa Indonesia." if language == "id" else "Answer concisely in English."
    instruction = (
        "Anda adalah asisten hukum. Ringkas dokumen berikut menjadi poin-poin yang jelas, sertakan subjudul jika relevan, dan sorot risiko/isu hukum."
        if language == "id"
        else "You are a legal assistant. Summarize the document into clear bullet points, add subheadings if relevant, and highlight legal risks/issues. "
    )
    prompt = f"{instruction} {lang_line}\n\nDokumen:\n\n{snippet}"
    summary = call_claude(prompt, max_tokens=800, temperature=0.2)
    return {"summary": summary}


@app.post("/ask")
def ask(req: AskRequest):
    session = SESSIONS.get(req.session_id)
    if not session:
        raise HTTPException(status_code=404, detail="Session tidak ditemukan")

    doc_text = session.get("text", "")
    if not doc_text:
        raise HTTPException(status_code=400, detail="Tidak ada teks dalam sesi")

    question = req.question.strip()
    if not question:
        raise HTTPException(status_code=400, detail="Pertanyaan tidak boleh kosong")

    snippet = doc_text[:15000]
    language = (req.lang or "id").lower()
    intro = (
        "Anda adalah asisten QnA yang hanya menjawab berdasarkan isi dokumen. Jika jawaban tidak ada, katakan tidak ditemukan. Jawab ringkas dalam bahasa Indonesia, dan bila relevan sertakan kutipan singkat dari dokumen.\n\n"
        if language == "id"
        else "You are a QnA assistant that only answers based on the document. If the answer is not present, say it is not found. Answer concisely in English and include brief quotes from the document when relevant.\n\n"
    )
    prompt = f"{intro}Dokumen:\n\n{snippet}\n\nPertanyaan/Question: {question}"
    answer = call_claude(prompt, max_tokens=800, temperature=0.0)
    return {"answer": answer}


# Level 2: Multi-file analyze with optional confidential mode and chat history
@app.post("/analyze")
async def analyze(
    files: Optional[List[UploadFile]] = File(None),
    text: Optional[str] = Form(None),
    confidential: Optional[str] = Form(None),
    preset: Optional[str] = Form(None),  # e.g., summary|risk|clauses|timeline
    lang: Optional[str] = Form(None),  # 'id' or 'en'
    db: Session = Depends(get_db),
):
    combined_text_parts: List[str] = []

    # Process files
    if files:
        for f in files:
            content = await f.read()
            if not content:
                continue
            name_lower = (f.filename or "").lower()
            ctype = (f.content_type or "").lower()
            is_pdf = "pdf" in ctype or name_lower.endswith(".pdf")
            is_txt = "text" in ctype or name_lower.endswith(".txt")
            if is_pdf:
                combined_text_parts.append(extract_text_from_pdf(content))
            elif is_txt:
                try:
                    combined_text_parts.append(content.decode("utf-8", errors="ignore"))
                except Exception:
                    combined_text_parts.append(content.decode("latin-1", errors="ignore"))

    if text and text.strip():
        combined_text_parts.append(text.strip())

    combined_text = "\n\n---\n\n".join([p for p in combined_text_parts if p]).strip()
    if not combined_text:
        raise HTTPException(status_code=400, detail="Tidak ada input untuk dianalisa")

    language = (lang or "id").lower()
    mode = (preset or "summary").lower()

    if language == "en":
        assistant_intro = "You are a legal assistant."
        lang_intro = "Answer in English."
        instruction_map = {
            "summary": "Summarize the document and list the key points.",
            "risk": "Identify legal risks/issues and provide recommended actions.",
            "clauses": "Extract important clauses and entities (parties, dates, amounts, articles).",
            "timeline": "Create a timeline of key events or obligations.",
        }
        citations_line = "Include 'citations' as brief quotes from the document that support the answer."
        schema_line = (
            "Return the output in JSON with the schema: {summary?: string, risks?: string[], recommendations?: string[], clauses?: string[], entities?: object, timeline?: string[], citations?: string[]}."
        )
        output_only = "Only output valid JSON."
        document_label = "Document"
    else:
        assistant_intro = "Anda adalah asisten hukum."
        lang_intro = "Jawab dalam bahasa Indonesia."
        instruction_map = {
            "summary": "Ringkas dokumen dan beri poin-poin utama.",
            "risk": "Identifikasi risiko/isu hukum dan rekomendasi langkah.",
            "clauses": "Ekstrak klausul & entitas penting (pihak, tanggal, nilai, pasal).",
            "timeline": "Buat garis waktu kejadian atau kewajiban utama.",
        }
        citations_line = "Sertakan 'citations' berupa kutipan singkat dari dokumen yang menjadi dasar jawaban."
        schema_line = (
            "Kembalikan output dalam JSON dengan schema: {summary?: string, risks?: string[], recommendations?: string[], clauses?: string[], entities?: object, timeline?: string[], citations?: string[]}."
        )
        output_only = "Output hanya JSON valid."
        document_label = "Dokumen"

    instruction = instruction_map.get(mode, instruction_map["summary"])
    prompt = (
        f"{assistant_intro} {instruction} {citations_line} {schema_line} {lang_intro}\n\n"
        f"{document_label}:\n{combined_text[:15000]}\n\n{output_only}"
    )
    raw = call_claude(prompt, max_tokens=900, temperature=0.2)
    import json
    parsed: dict = {}
    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = {"summary": raw}

    is_confidential = (confidential or "").strip() in {"1", "true", "True", "on"}

    chat_id: Optional[str] = None
    assistant_message_id: Optional[str] = None
    if not is_confidential:
        chat = Chat(title="Analisa Dokumen", confidential=False)
        db.add(chat)
        db.flush()
        chat_id = chat.id
        user_msg = Message(chat_id=chat.id, role="user", content=f"[{mode}/{language}]\n" + combined_text[:15000])
        db.add(user_msg)
        assistant_msg = Message(chat_id=chat.id, role="assistant", content=(parsed.get("summary") or raw))
        db.add(assistant_msg)
        db.flush()
        assistant_message_id = assistant_msg.id
        # Save file names metadata
        if files:
            for f in files:
                if f and f.filename:
                    db.add(FileRec(chat_id=chat.id, name=f.filename))
        chat.updated_at = datetime.utcnow()
        db.commit()

    return {"result": parsed.get("summary") or parsed, "details": parsed, "chat_id": chat_id, "assistant_message_id": assistant_message_id}


@app.post("/draft")
def draft(req: DraftRequest):
    language = (req.lang or "id").lower()
    tone = (req.tone or ("formal" if language == "id" else "formal")).lower()
    length = (req.length or "medium").lower()
    doc_type = (req.doc_type or "Dokumen").strip()
    reqs = (req.requirements or "").strip()

    # Use session text as context if provided
    ctx_text = ""
    if req.session_id:
        session = SESSIONS.get(req.session_id)
        if session:
            ctx_text = session.get("text", "")[:12000]

    if language == "en":
        header = (
            f"You are a legal drafting assistant. Write a {doc_type} in a {tone} tone. Aim for {length} length. "
            "Use clear headings and numbered clauses. Insert placeholders where information is missing."
        )
        req_label = "Requirements"
        ctx_label = "Context (optional)"
        out_hint = "Return well-structured, copy-ready text in English."
    else:
        header = (
            f"Anda adalah asisten penyusun dokumen hukum. Tulis {doc_type} dengan gaya {tone}. Panjang {length}. "
            "Gunakan heading yang jelas dan klausul bernomor. Gunakan placeholder jika data belum tersedia."
        )
        req_label = "Kebutuhan"
        ctx_label = "Konteks (opsional)"
        out_hint = "Kembalikan teks rapi siap salin dalam bahasa Indonesia."

    prompt = (
        f"{header}\n\n{req_label}:\n{reqs}\n\n{ctx_label}:\n{ctx_text}\n\n{out_hint}"
    )

    draft_text = call_claude(prompt, max_tokens=1200, temperature=0.2)
    return {"draft": draft_text}

@app.get("/chats")
def list_chats(db: Session = Depends(get_db)):
    chats = db.query(Chat).order_by(Chat.updated_at.desc()).limit(50).all()
    return {
        "chats": [
            {
                "id": c.id,
                "title": c.title,
                "confidential": c.confidential,
                "created_at": c.created_at.isoformat(),
                "updated_at": c.updated_at.isoformat(),
            }
            for c in chats
        ]
    }


@app.post("/export_pdf")
def export_pdf(chat_id: str = Form(...), db: Session = Depends(get_db)):
    chat = db.query(Chat).filter(Chat.id == chat_id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat tidak ditemukan")

    messages = db.query(Message).filter(Message.chat_id == chat_id).order_by(Message.created_at.asc()).all()
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    title = chat.title or "Analisa Dokumen"
    pdf.multi_cell(0, 10, txt=title)
    pdf.ln(5)
    for m in messages:
        header = f"[{m.role.upper()}] {m.created_at.isoformat()}\n"
        pdf.set_font("Arial", style="B", size=11)
        pdf.multi_cell(0, 8, txt=header)
        pdf.set_font("Arial", size=11)
        pdf.multi_cell(0, 6, txt=m.content)
        pdf.ln(2)

    pdf_bytes = pdf.output(dest="S").encode("latin-1", errors="ignore")
    from fastapi.responses import Response

    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=chat_{chat_id}.pdf"})


@app.post("/export_docx")
def export_docx(chat_id: str = Form(...), db: Session = Depends(get_db)):
    from docx import Document
    from docx.shared import Pt
    chat = db.query(Chat).filter(Chat.id == chat_id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat tidak ditemukan")
    messages = db.query(Message).filter(Message.chat_id == chat_id).order_by(Message.created_at.asc()).all()
    doc = Document()
    doc.add_heading(chat.title or "Analisa Dokumen", level=1)
    for m in messages:
        p = doc.add_paragraph()
        run = p.add_run(f"[{m.role.upper()}] {m.created_at.isoformat()}\n")
        run.bold = True
        p2 = doc.add_paragraph(m.content)
        p2_format = p2.paragraph_format
    from fastapi.responses import Response
    b = bytes()
    import io
    buf = io.BytesIO()
    doc.save(buf)
    b = buf.getvalue()
    return Response(content=b, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=chat_{chat_id}.docx"})


@app.post("/feedback")
def feedback(chat_id: Optional[str] = Form(None), message_id: Optional[str] = Form(None), value: int = Form(...), comment: Optional[str] = Form(None), db: Session = Depends(get_db)):
    if value not in (-1, 1):
        raise HTTPException(status_code=400, detail="value harus -1 atau 1")
    fb = Feedback(chat_id=chat_id, message_id=message_id, value=value, comment=comment)
    db.add(fb)
    db.commit()
    return {"ok": True}


# Draft export endpoints
@app.post("/export_draft_pdf")
def export_draft_pdf(text: str = Form(...), title: Optional[str] = Form("Draft Dokumen")):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", style="B", size=14)
    pdf.multi_cell(0, 10, txt=title or "Draft")
    pdf.ln(4)
    pdf.set_font("Arial", size=11)
    for line in (text or "").split("\n"):
        pdf.multi_cell(0, 6, txt=line)
    pdf_bytes = pdf.output(dest="S").encode("latin-1", errors="ignore")
    # Use inline so the browser opens a preview tab; users can download from there
    return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": "inline; filename=draft.pdf"})


@app.post("/export_draft_docx")
def export_draft_docx(text: str = Form(...), title: Optional[str] = Form("Draft Dokumen")):
    from docx import Document
    doc = Document()
    doc.add_heading(title or "Draft", level=1)
    for para in (text or "").split("\n\n"):
        p = doc.add_paragraph()
        for line in para.split("\n"):
            p.add_run(line)
            p.add_run("\n")
    import io
    buf = io.BytesIO()
    doc.save(buf)
    b = buf.getvalue()
    return Response(content=b, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": "attachment; filename=draft.docx"})


# Document review for submission readiness and change summary
@app.post("/review")
async def review(
    file_current: Optional[UploadFile] = File(None),
    text_current: Optional[str] = Form(None),
    file_previous: Optional[UploadFile] = File(None),
    text_previous: Optional[str] = Form(None),
    lang: Optional[str] = Form(None),
):
    def read_any(file: Optional[UploadFile], text: Optional[str]) -> str:
        if text and text.strip():
            return text.strip()
        if file:
            data = await_read(file)  # type: ignore
        else:
            return ""
        return data

    # helper for sync/async read
    async def read_upload(file: UploadFile) -> bytes:
        return await file.read()

    def decode_bytes(name: str, ctype: str, data: bytes) -> str:
        name_lower = (name or "").lower()
        ctype_lower = (ctype or "").lower()
        if "pdf" in ctype_lower or name_lower.endswith(".pdf"):
            return extract_text_from_pdf(data)
        try:
            return data.decode("utf-8", errors="ignore")
        except Exception:
            return data.decode("latin-1", errors="ignore")

    # read current
    current_text = (text_current or "").strip()
    if not current_text and file_current is not None:
        data = await file_current.read()
        current_text = decode_bytes(file_current.filename or "", file_current.content_type or "", data)

    # read previous if any
    previous_text = (text_previous or "").strip()
    if not previous_text and file_previous is not None:
        data_p = await file_previous.read()
        previous_text = decode_bytes(file_previous.filename or "", file_previous.content_type or "", data_p)

    if not current_text:
        raise HTTPException(status_code=400, detail="Tidak ada dokumen untuk direview")

    language = (lang or "id").lower()
    if language == "en":
        intro = (
            "You are a legal reviewer. Assess the document's readiness for court submission (formalities, identities, references, signatures, annexes). "
            "Identify gaps or issues, and provide clear recommendations."
        )
        change_part = (
            "Also summarize substantive changes between CURRENT and PREVIOUS if PREVIOUS is provided."
            if previous_text
            else ""
        )
        schema_hint = (
            "Return JSON: {summary: string, missing?: string[], issues?: string[], changes?: string[], recommendations?: string[], citations?: string[]}"
        )
        labels = ("CURRENT", "PREVIOUS", "Document")
    else:
        intro = (
            "Anda adalah reviewer dokumen. Nilai kesiapan dokumen untuk pengajuan ke pengadilan (formalitas, identitas, rujukan, tanda tangan, lampiran). "
            "Identifikasi kekurangan/isu dan beri rekomendasi yang jelas."
        )
        change_part = (
            "Ringkas perubahan substansial antara DOKUMEN SAAT INI dan DOKUMEN SEBELUMNYA jika tersedia."
            if previous_text
            else ""
        )
        schema_hint = (
            "Kembalikan JSON: {summary: string, missing?: string[], issues?: string[], changes?: string[], recommendations?: string[], citations?: string[]}"
        )
        labels = ("DOKUMEN SAAT INI", "DOKUMEN SEBELUMNYA", "Dokumen")

    prompt = (
        f"{intro} {change_part} {schema_hint}\n\n"
        f"{labels[0]}:\n{current_text[:15000]}\n\n"
        + (f"{labels[1]}:\n{previous_text[:8000]}\n\n" if previous_text else "")
        + ("Output hanya JSON valid." if language == "id" else "Only output valid JSON.")
    )

    raw = call_claude(prompt, max_tokens=900, temperature=0.2)
    import json
    try:
        parsed = json.loads(raw)
    except Exception:
        parsed = {"summary": raw}

    return {"review": parsed}


@app.post("/compare")
async def compare(file_a: Optional[UploadFile] = File(None), file_b: Optional[UploadFile] = File(None), text_a: Optional[str] = Form(None), text_b: Optional[str] = Form(None)):
    if not ((file_a or text_a) and (file_b or text_b)):
        raise HTTPException(status_code=400, detail="Butuh dua dokumen untuk dibandingkan")
    def read_any(file: Optional[UploadFile], text: Optional[str]) -> str:
        if text and text.strip():
            return text.strip()
        if file:
            data = file.file.read()
            ctype = (file.content_type or "").lower()
            name = (file.filename or "").lower()
            if "pdf" in ctype or name.endswith(".pdf"):
                return extract_text_from_pdf(data)
            try:
                return data.decode("utf-8", errors="ignore")
            except Exception:
                return data.decode("latin-1", errors="ignore")
        return ""
    a = read_any(file_a, text_a)
    b = read_any(file_b, text_b)
    prompt = (
        "Bandingkan dua dokumen kontrak berikut. Ringkas perbedaan utama, risiko akibat perubahan, dan bagian yang identik. "
        "Tampilkan sebagai poin-poin. Jawab dalam bahasa Indonesia.\n\nDokumen A:\n" + a[:8000] + "\n\nDokumen B:\n" + b[:8000]
    )
    out = call_claude(prompt, max_tokens=800, temperature=0.2)
    return {"diff": out}


@app.post("/compliance")
async def compliance(file_doc: Optional[UploadFile] = File(None), file_template: Optional[UploadFile] = File(None), text_doc: Optional[str] = Form(None), text_template: Optional[str] = Form(None)):
    def read_any(file: Optional[UploadFile], text: Optional[str]) -> str:
        if text and text.strip():
            return text.strip()
        if file:
            data = file.file.read()
            ctype = (file.content_type or "").lower()
            name = (file.filename or "").lower()
            if "pdf" in ctype or name.endswith(".pdf"):
                return extract_text_from_pdf(data)
            try:
                return data.decode("utf-8", errors="ignore")
            except Exception:
                return data.decode("latin-1", errors="ignore")
        return ""
    doc = read_any(file_doc, text_doc)
    tpl = read_any(file_template, text_template)
    if not (doc and tpl):
        raise HTTPException(status_code=400, detail="Butuh dokumen dan template")
    prompt = (
        "Cek kepatuhan dokumen terhadap template standar. Daftarkan: bagian sesuai, deviasi, dan kekurangan yang perlu ditambah. "
        "Berikan rekomendasi singkat. Jawab dalam bahasa Indonesia.\n\nTemplate:\n" + tpl[:8000] + "\n\nDokumen:\n" + doc[:8000]
    )
    out = call_claude(prompt, max_tokens=900, temperature=0.2)
    return {"compliance": out}


# Local dev entrypoint (optional)
if __name__ == "__main__":  # pragma: no cover
    import uvicorn

    port = int(os.getenv("PORT", "8000"))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=True)


